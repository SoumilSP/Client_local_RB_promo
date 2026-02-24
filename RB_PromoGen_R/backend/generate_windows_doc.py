from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# Set document title
title = doc.add_heading('Trade Promotion Optimizer', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Technical Documentation & Windows Server Deployment Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(196, 25, 116)

date_para = doc.add_paragraph(f'Version: 2.1 | Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(10)
date_para.runs[0].font.italic = True

doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Architecture Overview',
    '3. R Engine Files - Detailed Description',
    '4. Data Upload & Processing Flow',
    '5. Frontend React Components',
    '6. API Endpoints & Data Flow',
    '7. Optimization Calculation Logic',
    '8. Configuration Parameters',
    '9. Output & Results Rendering',
    '10. File Directory Structure',
    '11. Windows Server Production Deployment',
    '12. Troubleshooting Guide',
    '13. Appendix'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Section 1: Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The Trade Promotion Optimizer is a sophisticated FMCG promotional planning tool '
    'that combines a React frontend with an R-powered analytics engine. The system '
    'enables category managers to create optimized promotional calendars by processing '
    'real data from uploaded Excel files through the R Plumber API.'
)

doc.add_heading('Technology Stack', level=2)
tech_items = [
    'Frontend: React 18.x + Tailwind CSS + shadcn/ui components',
    'Backend Proxy: Python FastAPI (handles file uploads, proxies to R)',
    'Analytics Engine: R 4.2.2 with Plumber REST API',
    'Data Storage: File-based system (configurable directory)',
    'Charts: Recharts library for data visualization'
]
for item in tech_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 2-10 (abbreviated for brevity - same as before)
doc.add_heading('2. System Architecture Overview', level=1)
doc.add_paragraph('Please refer to the full technical documentation for detailed architecture diagrams.')

doc.add_heading('3-10. Technical Details', level=1)
doc.add_paragraph(
    'Sections 3-10 contain detailed technical information about:\n'
    '• R Engine Files and their purposes\n'
    '• Data Upload & Processing Flow\n'
    '• Frontend React Components\n'
    '• API Endpoints & Data Flow\n'
    '• Optimization Calculation Logic\n'
    '• Configuration Parameters\n'
    '• Output & Results Rendering\n'
    '• File Directory Structure\n\n'
    'Please refer to the main technical document for these sections.'
)

doc.add_page_break()

# Section 11: Windows Server Production Deployment
doc.add_heading('11. Windows Server Production Deployment', level=1)

doc.add_heading('11.1 Prerequisites', level=2)

prereq_table = doc.add_table(rows=7, cols=4)
prereq_table.style = 'Table Grid'

prereq_headers = ['Software', 'Version', 'Download URL', 'Purpose']
for i, h in enumerate(prereq_headers):
    prereq_table.rows[0].cells[i].text = h

prereq_data = [
    ('Windows Server', '2019/2022', 'Microsoft Volume Licensing', 'Operating System'),
    ('Node.js', '18.x LTS or 20.x LTS', 'https://nodejs.org/', 'React frontend runtime'),
    ('Python', '3.10 or 3.11', 'https://python.org/downloads/', 'FastAPI backend runtime'),
    ('R', '4.2.x or 4.3.x', 'https://cran.r-project.org/', 'Analytics engine'),
    ('Git', 'Latest', 'https://git-scm.com/', 'Source code management'),
    ('NSSM', 'Latest', 'https://nssm.cc/', 'Windows Service wrapper')
]

for i, (soft, ver, url, purpose) in enumerate(prereq_data, 1):
    prereq_table.rows[i].cells[0].text = soft
    prereq_table.rows[i].cells[1].text = ver
    prereq_table.rows[i].cells[2].text = url
    prereq_table.rows[i].cells[3].text = purpose

doc.add_paragraph()

doc.add_heading('11.2 Step 1: Install Prerequisites', level=2)

install_steps = '''
A. INSTALL NODE.JS
------------------
1. Download Node.js LTS from https://nodejs.org/
2. Run installer with default options
3. Verify installation:
   > node --version
   > npm --version

B. INSTALL PYTHON
-----------------
1. Download Python 3.11 from https://python.org/downloads/
2. Run installer - IMPORTANT: Check "Add Python to PATH"
3. Verify installation:
   > python --version
   > pip --version

C. INSTALL R
------------
1. Download R from https://cran.r-project.org/bin/windows/base/
2. Run installer with default options
3. Add R to system PATH:
   - Right-click "This PC" → Properties → Advanced system settings
   - Click "Environment Variables"
   - Edit "Path" under System variables
   - Add: C:\\Program Files\\R\\R-4.3.x\\bin
4. Verify installation:
   > Rscript --version

D. INSTALL R PACKAGES
---------------------
Open R console (or RStudio) and run:
> install.packages(c("plumber", "jsonlite", "data.table", "readxl", "dplyr", "tidyr"))

E. INSTALL GIT
--------------
1. Download from https://git-scm.com/download/win
2. Run installer with default options
3. Verify: > git --version

F. INSTALL NSSM (Non-Sucking Service Manager)
---------------------------------------------
1. Download from https://nssm.cc/download
2. Extract to C:\\nssm
3. Add C:\\nssm\\win64 to system PATH
'''

doc.add_paragraph(install_steps)

doc.add_page_break()

doc.add_heading('11.3 Step 2: Clone and Setup Application', level=2)

setup_steps = '''
A. CREATE DIRECTORY STRUCTURE
-----------------------------
> mkdir C:\\TradePromoOptimizer
> cd C:\\TradePromoOptimizer
> mkdir frontend backend data logs

B. CLONE OR COPY SOURCE CODE
----------------------------
Option 1 - Git Clone (if repository available):
> git clone <repository-url> .

Option 2 - Copy Files:
Copy the following directories from development environment:
- /app/frontend → C:\\TradePromoOptimizer\\frontend
- /app/backend → C:\\TradePromoOptimizer\\backend

C. SETUP FRONTEND
-----------------
> cd C:\\TradePromoOptimizer\\frontend
> npm install
> npm run build

This creates a production build in C:\\TradePromoOptimizer\\frontend\\build

D. SETUP BACKEND
----------------
> cd C:\\TradePromoOptimizer\\backend
> pip install -r requirements.txt

If requirements.txt doesn't exist, install manually:
> pip install fastapi uvicorn httpx python-multipart pydantic

E. CREATE DATA DIRECTORY
------------------------
> mkdir C:\\TradePromoOptimizer\\data

This directory will store:
- Uploaded Excel/CSV files
- file_mapping.json
'''

doc.add_paragraph(setup_steps)

doc.add_page_break()

doc.add_heading('11.4 Step 3: Configure Environment Variables', level=2)

env_config = '''
A. FRONTEND CONFIGURATION
-------------------------
Create/Edit: C:\\TradePromoOptimizer\\frontend\\.env.production

Contents:
REACT_APP_BACKEND_URL=http://your-server-ip:8001

For production with domain:
REACT_APP_BACKEND_URL=https://promo-optimizer.yourcompany.com

B. BACKEND CONFIGURATION
------------------------
Create/Edit: C:\\TradePromoOptimizer\\backend\\.env

Contents:
R_API_URL=http://localhost:8002
DATA_DIR=C:\\TradePromoOptimizer\\data
CORS_ORIGINS=http://your-server-ip:3000,https://promo-optimizer.yourcompany.com
LOG_LEVEL=INFO

C. R ENGINE CONFIGURATION
-------------------------
Edit: C:\\TradePromoOptimizer\\backend\\r_engine\\plumber_api.R

Update line 13:
DATA_DIR <- "C:/TradePromoOptimizer/data"

Note: Use forward slashes (/) in R paths, even on Windows

D. WINDOWS ENVIRONMENT VARIABLES (System-wide)
----------------------------------------------
1. Open System Properties → Environment Variables
2. Add new System variables:
   - TRADE_PROMO_HOME = C:\\TradePromoOptimizer
   - R_HOME = C:\\Program Files\\R\\R-4.3.x
'''

doc.add_paragraph(env_config)

doc.add_page_break()

doc.add_heading('11.5 Step 4: Create Windows Services', level=2)

services_config = '''
Use NSSM to create Windows services for each component:

A. CREATE R ENGINE SERVICE
--------------------------
Open Command Prompt as Administrator:

> nssm install TradePromoR

In the NSSM GUI:
- Path: C:\\Program Files\\R\\R-4.3.x\\bin\\Rscript.exe
- Startup directory: C:\\TradePromoOptimizer\\backend\\r_engine
- Arguments: start_api.R
- Service name: TradePromoR

Click "Install service"

Configure service:
> nssm set TradePromoR DisplayName "Trade Promo Optimizer - R Engine"
> nssm set TradePromoR Description "R Plumber API for Trade Promotion Analytics"
> nssm set TradePromoR Start SERVICE_AUTO_START
> nssm set TradePromoR AppStdout C:\\TradePromoOptimizer\\logs\\r_engine.log
> nssm set TradePromoR AppStderr C:\\TradePromoOptimizer\\logs\\r_engine_error.log

B. CREATE BACKEND SERVICE
-------------------------
> nssm install TradePromoBackend

In the NSSM GUI:
- Path: C:\\Python311\\python.exe (adjust to your Python path)
- Startup directory: C:\\TradePromoOptimizer\\backend
- Arguments: -m uvicorn server:app --host 0.0.0.0 --port 8001
- Service name: TradePromoBackend

Configure:
> nssm set TradePromoBackend DisplayName "Trade Promo Optimizer - Backend API"
> nssm set TradePromoBackend Description "FastAPI backend for Trade Promotion Optimizer"
> nssm set TradePromoBackend Start SERVICE_AUTO_START
> nssm set TradePromoBackend AppStdout C:\\TradePromoOptimizer\\logs\\backend.log
> nssm set TradePromoBackend AppStderr C:\\TradePromoOptimizer\\logs\\backend_error.log
> nssm set TradePromoBackend DependOnService TradePromoR

C. CREATE FRONTEND SERVICE (Optional - for Node.js serve)
---------------------------------------------------------
Install serve globally:
> npm install -g serve

Create service:
> nssm install TradePromoFrontend

In the NSSM GUI:
- Path: C:\\Users\\<user>\\AppData\\Roaming\\npm\\serve.cmd
- Startup directory: C:\\TradePromoOptimizer\\frontend\\build
- Arguments: -s . -l 3000
- Service name: TradePromoFrontend

Configure:
> nssm set TradePromoFrontend DisplayName "Trade Promo Optimizer - Frontend"
> nssm set TradePromoFrontend Start SERVICE_AUTO_START
> nssm set TradePromoFrontend DependOnService TradePromoBackend

D. START ALL SERVICES
---------------------
> net start TradePromoR
> net start TradePromoBackend
> net start TradePromoFrontend

Or use Services console (services.msc) to start/stop services.
'''

doc.add_paragraph(services_config)

doc.add_page_break()

doc.add_heading('11.6 Step 5: Configure IIS as Reverse Proxy (Recommended)', level=2)

iis_config = '''
Using IIS provides better security, SSL termination, and enterprise features.

A. INSTALL IIS WITH REQUIRED FEATURES
-------------------------------------
1. Open Server Manager → Add Roles and Features
2. Select "Web Server (IIS)"
3. Include these features:
   - Application Development → WebSocket Protocol
   - Health and Diagnostics → HTTP Logging
   - Security → Request Filtering
   - Management Tools → IIS Management Console

B. INSTALL URL REWRITE & ARR
----------------------------
1. Download URL Rewrite: https://www.iis.net/downloads/microsoft/url-rewrite
2. Download Application Request Routing (ARR): 
   https://www.iis.net/downloads/microsoft/application-request-routing
3. Install both MSI packages

C. ENABLE ARR PROXY
-------------------
1. Open IIS Manager
2. Select server node (top level)
3. Double-click "Application Request Routing Cache"
4. Click "Server Proxy Settings"
5. Check "Enable proxy"
6. Click Apply

D. CREATE IIS WEBSITE
---------------------
1. In IIS Manager, right-click "Sites" → "Add Website"
2. Site name: TradePromoOptimizer
3. Physical path: C:\\TradePromoOptimizer\\frontend\\build
4. Binding: 
   - Type: http (or https with SSL certificate)
   - Port: 80 (or 443 for https)
   - Host name: promo-optimizer.yourcompany.com

E. CONFIGURE URL REWRITE RULES
------------------------------
1. Select the new website
2. Double-click "URL Rewrite"
3. Click "Add Rule(s)..." → "Blank rule"

Rule 1 - API Proxy:
- Name: API Proxy to Backend
- Match URL:
  - Requested URL: Matches the Pattern
  - Using: Wildcards
  - Pattern: api/*
- Action:
  - Action type: Rewrite
  - Rewrite URL: http://localhost:8001/{R:0}

Rule 2 - R Engine Proxy:
- Name: R Engine Proxy
- Match URL:
  - Pattern: r/*
- Action:
  - Rewrite URL: http://localhost:8002/{R:1}

F. WEB.CONFIG FOR SPA ROUTING
-----------------------------
Create C:\\TradePromoOptimizer\\frontend\\build\\web.config:

<?xml version="1.0" encoding="UTF-8"?>
<configuration>
  <system.webServer>
    <rewrite>
      <rules>
        <rule name="API Proxy" stopProcessing="true">
          <match url="^api/(.*)" />
          <action type="Rewrite" url="http://localhost:8001/api/{R:1}" />
        </rule>
        <rule name="React SPA" stopProcessing="true">
          <match url=".*" />
          <conditions logicalGrouping="MatchAll">
            <add input="{REQUEST_FILENAME}" matchType="IsFile" negate="true" />
            <add input="{REQUEST_FILENAME}" matchType="IsDirectory" negate="true" />
          </conditions>
          <action type="Rewrite" url="/" />
        </rule>
      </rules>
    </rewrite>
    <staticContent>
      <mimeMap fileExtension=".json" mimeType="application/json" />
    </staticContent>
  </system.webServer>
</configuration>
'''

doc.add_paragraph(iis_config)

doc.add_page_break()

doc.add_heading('11.7 Step 6: Configure Firewall', level=2)

firewall_config = '''
A. WINDOWS FIREWALL RULES
-------------------------
Open Windows Defender Firewall with Advanced Security:

1. Inbound Rule for HTTP (if not using IIS):
   - Name: Trade Promo Optimizer - HTTP
   - Port: 80, 3000
   - Action: Allow

2. Inbound Rule for HTTPS:
   - Name: Trade Promo Optimizer - HTTPS
   - Port: 443
   - Action: Allow

3. Internal ports (should NOT be exposed externally):
   - Port 8001 (Backend API) - localhost only
   - Port 8002 (R Engine) - localhost only

PowerShell commands:
> New-NetFirewallRule -DisplayName "Trade Promo HTTP" -Direction Inbound -Port 80 -Protocol TCP -Action Allow
> New-NetFirewallRule -DisplayName "Trade Promo HTTPS" -Direction Inbound -Port 443 -Protocol TCP -Action Allow

B. CORPORATE FIREWALL
---------------------
Request the following from network team:
- Allow inbound HTTP/HTTPS to server IP
- Allow outbound for Windows Updates and package downloads (npm, pip, CRAN)
'''

doc.add_paragraph(firewall_config)

doc.add_page_break()

doc.add_heading('11.8 Step 7: SSL/HTTPS Configuration', level=2)

ssl_config = '''
A. OBTAIN SSL CERTIFICATE
-------------------------
Option 1 - Corporate CA:
- Request certificate from your IT security team
- Provide: Common Name (CN) = promo-optimizer.yourcompany.com
- Export as PFX file with private key

Option 2 - Let's Encrypt (for public-facing servers):
- Install win-acme: https://www.win-acme.com/
- Run: > wacs.exe
- Follow prompts to generate certificate

Option 3 - Self-Signed (Development/Testing only):
PowerShell:
> New-SelfSignedCertificate -DnsName "promo-optimizer.local" -CertStoreLocation cert:\\LocalMachine\\My

B. INSTALL CERTIFICATE IN IIS
-----------------------------
1. Open IIS Manager
2. Select server node → "Server Certificates"
3. Click "Import" (for PFX) or "Complete Certificate Request" (for CA-signed)
4. Select the website → Bindings
5. Add binding:
   - Type: https
   - Port: 443
   - SSL certificate: Select your certificate

C. FORCE HTTPS REDIRECT
-----------------------
Add to web.config:

<rule name="HTTPS Redirect" stopProcessing="true">
  <match url="(.*)" />
  <conditions>
    <add input="{HTTPS}" pattern="off" ignoreCase="true" />
  </conditions>
  <action type="Redirect" url="https://{HTTP_HOST}/{R:1}" redirectType="Permanent" />
</rule>
'''

doc.add_paragraph(ssl_config)

doc.add_page_break()

doc.add_heading('11.9 Step 8: Production Monitoring & Maintenance', level=2)

monitoring_config = '''
A. LOG FILE LOCATIONS
---------------------
- R Engine: C:\\TradePromoOptimizer\\logs\\r_engine.log
- Backend: C:\\TradePromoOptimizer\\logs\\backend.log
- IIS: C:\\inetpub\\logs\\LogFiles\\W3SVC<site-id>\\

B. HEALTH CHECK SCRIPT
----------------------
Create: C:\\TradePromoOptimizer\\scripts\\health_check.ps1

$services = @("TradePromoR", "TradePromoBackend", "TradePromoFrontend")
$healthEndpoint = "http://localhost:8001/api/health"

foreach ($service in $services) {
    $status = Get-Service -Name $service -ErrorAction SilentlyContinue
    if ($status.Status -ne "Running") {
        Write-Host "WARNING: $service is not running" -ForegroundColor Red
        Start-Service -Name $service
    }
}

try {
    $response = Invoke-RestMethod -Uri $healthEndpoint -TimeoutSec 10
    Write-Host "API Health: OK" -ForegroundColor Green
} catch {
    Write-Host "API Health: FAILED" -ForegroundColor Red
}

Schedule with Task Scheduler to run every 5 minutes.

C. BACKUP SCRIPT
----------------
Create: C:\\TradePromoOptimizer\\scripts\\backup.ps1

$date = Get-Date -Format "yyyyMMdd"
$backupDir = "C:\\Backups\\TradePromo\\$date"
New-Item -ItemType Directory -Path $backupDir -Force

# Backup data files
Copy-Item -Path "C:\\TradePromoOptimizer\\data\\*" -Destination "$backupDir\\data" -Recurse

# Backup configuration
Copy-Item -Path "C:\\TradePromoOptimizer\\backend\\.env" -Destination "$backupDir\\"
Copy-Item -Path "C:\\TradePromoOptimizer\\frontend\\.env.production" -Destination "$backupDir\\"

Write-Host "Backup completed to $backupDir"

Schedule daily via Task Scheduler.

D. LOG ROTATION
---------------
Use PowerShell to rotate logs weekly:

Get-ChildItem "C:\\TradePromoOptimizer\\logs\\*.log" | 
Where-Object { $_.LastWriteTime -lt (Get-Date).AddDays(-7) } | 
ForEach-Object { 
    Compress-Archive -Path $_.FullName -DestinationPath "$($_.FullName).zip"
    Remove-Item $_.FullName 
}

E. WINDOWS EVENT LOG
--------------------
Service issues are logged to Windows Event Log:
- Open Event Viewer
- Navigate to: Windows Logs → Application
- Filter by Source: nssm
'''

doc.add_paragraph(monitoring_config)

doc.add_page_break()

doc.add_heading('11.10 Production Checklist', level=2)

checklist = '''
PRE-DEPLOYMENT CHECKLIST
========================

[ ] Windows Server 2019/2022 installed and updated
[ ] Node.js 18+ LTS installed
[ ] Python 3.10+ installed with pip
[ ] R 4.2+ installed with required packages
[ ] NSSM installed and in PATH
[ ] IIS installed with URL Rewrite and ARR

[ ] Application files copied to C:\\TradePromoOptimizer
[ ] Frontend build completed (npm run build)
[ ] Backend dependencies installed (pip install -r requirements.txt)

[ ] Environment variables configured:
    [ ] Frontend .env.production
    [ ] Backend .env
    [ ] R engine DATA_DIR path

[ ] Windows Services created:
    [ ] TradePromoR
    [ ] TradePromoBackend
    [ ] TradePromoFrontend (optional)

[ ] IIS configured:
    [ ] Website created
    [ ] URL Rewrite rules added
    [ ] SSL certificate installed
    [ ] HTTPS redirect enabled

[ ] Firewall rules configured
[ ] Health check script scheduled
[ ] Backup script scheduled
[ ] Log rotation configured

POST-DEPLOYMENT VERIFICATION
============================

1. Open browser to https://promo-optimizer.yourcompany.com
2. Verify R Engine status shows "Connected" (green)
3. Upload a test data file
4. Run optimization with default settings
5. Verify results display correctly
6. Check all log files for errors
'''

doc.add_paragraph(checklist)

doc.add_page_break()

doc.add_heading('11.11 Troubleshooting Windows Deployment', level=2)

win_troubleshooting = '''
COMMON ISSUES AND SOLUTIONS
===========================

Issue: R Engine service won't start
-----------------------------------
1. Check R is in PATH: > Rscript --version
2. Check logs: C:\\TradePromoOptimizer\\logs\\r_engine_error.log
3. Verify R packages installed:
   > R -e "library(plumber); library(jsonlite)"
4. Check port 8002 not in use: > netstat -ano | findstr :8002

Issue: Backend can't connect to R Engine
----------------------------------------
1. Verify R service is running: > sc query TradePromoR
2. Test R API directly: > curl http://localhost:8002/health
3. Check backend .env has correct R_API_URL
4. Verify no firewall blocking localhost connections

Issue: Frontend shows blank page
--------------------------------
1. Check frontend build exists in C:\\TradePromoOptimizer\\frontend\\build
2. Verify IIS website points to build directory
3. Check browser console (F12) for JavaScript errors
4. Ensure web.config exists with SPA routing rules

Issue: API calls return 404
---------------------------
1. Verify URL Rewrite rules in IIS
2. Check ARR proxy is enabled
3. Test backend directly: > curl http://localhost:8001/api/health
4. Check IIS logs for request routing

Issue: File upload fails
------------------------
1. Check IIS request size limits (default 30MB):
   In web.config: <requestLimits maxAllowedContentLength="104857600" />
2. Verify data directory permissions
3. Ensure IUSR has write access to C:\\TradePromoOptimizer\\data

Issue: Services crash after Windows Update
------------------------------------------
1. Re-register services with NSSM
2. Check system PATH variables still correct
3. Verify no port conflicts from new Windows services
4. Review Windows Event Log for service-specific errors

PERFORMANCE TUNING
==================

1. Increase IIS worker process memory limit
2. Configure R to use multiple cores:
   In plumber_api.R: options(mc.cores = 4)
3. Enable IIS output caching for static files
4. Consider separate server for R engine if heavy load
'''

doc.add_paragraph(win_troubleshooting)

doc.add_page_break()

# Section 12: Troubleshooting Guide (General)
doc.add_heading('12. Troubleshooting Guide (General)', level=1)

troubleshooting = [
    ('R Engine Not Starting', 
     'Check R installation: Rscript --version\n'
     'Install packages: R -e "install.packages(c(\'plumber\', \'jsonlite\', \'data.table\', \'readxl\'))"'),
    
    ('Data Not Loading in R',
     'Check file_mapping.json exists\n'
     'Verify file paths use forward slashes in R\n'
     'Check R logs for file read errors'),
    
    ('Optimization Returns Empty Results',
     'Ensure at least one mechanic is selected\n'
     'Check that uploaded files have data\n'
     'Verify constraints are not too restrictive')
]

for issue, solution in troubleshooting:
    doc.add_heading(issue, level=2)
    doc.add_paragraph(solution)

doc.add_page_break()

# Section 13: Appendix
doc.add_heading('13. Appendix', level=1)

doc.add_heading('13.1 Quick Reference - Windows Commands', level=2)

commands = '''
SERVICE MANAGEMENT:
> net start TradePromoR          # Start R engine
> net stop TradePromoR           # Stop R engine
> net start TradePromoBackend    # Start backend
> sc query TradePromoR           # Check service status
> nssm restart TradePromoR       # Restart via NSSM

LOG VIEWING:
> Get-Content C:\\TradePromoOptimizer\\logs\\r_engine.log -Tail 50
> Get-Content C:\\TradePromoOptimizer\\logs\\backend.log -Tail 50 -Wait

TESTING:
> curl http://localhost:8002/health           # Test R engine
> curl http://localhost:8001/api/health       # Test backend
> Invoke-RestMethod http://localhost:8001/api/health  # PowerShell

IIS:
> iisreset                        # Restart IIS
> appcmd list sites               # List websites
> appcmd recycle apppool /apppool.name:DefaultAppPool
'''

doc.add_paragraph(commands)

doc.add_heading('13.2 Port Reference', level=2)

port_table = doc.add_table(rows=5, cols=3)
port_table.style = 'Table Grid'

port_headers = ['Port', 'Service', 'Access']
for i, h in enumerate(port_headers):
    port_table.rows[0].cells[i].text = h

ports = [
    ('80/443', 'IIS (Frontend)', 'External - Users access here'),
    ('3000', 'Node serve (if not using IIS)', 'Internal only'),
    ('8001', 'FastAPI Backend', 'Internal only - proxied via IIS'),
    ('8002', 'R Plumber API', 'Internal only - called by backend')
]

for i, (port, service, access) in enumerate(ports, 1):
    port_table.rows[i].cells[0].text = port
    port_table.rows[i].cells[1].text = service
    port_table.rows[i].cells[2].text = access

# Save document
doc.save('/app/frontend/public/Trade_Promotion_Optimizer_Windows_Deployment.docx')
print('Windows Deployment document created successfully!')
