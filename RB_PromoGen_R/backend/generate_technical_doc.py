from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# Set document title
title = doc.add_heading('Trade Promotion Optimizer', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Technical Documentation & System Architecture')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(196, 25, 116)  # Reckitt pink

# Add date
date_para = doc.add_paragraph(f'Version: 2.0 | Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(10)
date_para.runs[0].font.italic = True

doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Architecture Overview',
    '3. R Engine Files - Detailed Description',
    '4. Data Upload & Processing Flow',
    '5. Frontend React Components',
    '6. API Endpoints & Data Flow',
    '7. Optimization Calculation Logic',
    '8. Configuration Parameters',
    '9. Output & Results Rendering',
    '10. File Directory Structure',
    '11. Troubleshooting Guide',
    '12. Appendix'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Section 1: Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The Trade Promotion Optimizer is a sophisticated FMCG promotional planning tool '
    'that combines a React frontend with an R-powered analytics engine. The system '
    'enables category managers to create optimized promotional calendars by processing '
    'real data from uploaded Excel files through the R Plumber API.'
)

doc.add_heading('Technology Stack', level=2)
tech_items = [
    'Frontend: React 18.x + Tailwind CSS + shadcn/ui components',
    'Backend Proxy: Python FastAPI (handles file uploads, proxies to R)',
    'Analytics Engine: R 4.2.2 with Plumber REST API',
    'Data Storage: File-based system (/app/data directory)',
    'Charts: Recharts library for data visualization'
]
for item in tech_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# Section 2: System Architecture
doc.add_heading('2. System Architecture Overview', level=1)

doc.add_heading('2.1 Three-Tier Architecture', level=2)
doc.add_paragraph(
    'The application follows a three-tier architecture with clear separation of concerns:'
)

arch_diagram = '''
┌─────────────────────────────────────────────────────────────────────┐
│                         USER INTERFACE                               │
│                    React Frontend (Port 3000)                        │
│   ┌─────────────┐  ┌─────────────┐  ┌─────────────────────────┐    │
│   │ DataInput   │  │ Optimizer   │  │ Other Pages             │    │
│   │ Page.js     │  │ NewPage.js  │  │ (PromoAnalysis, etc.)   │    │
│   └─────────────┘  └─────────────┘  └─────────────────────────┘    │
└───────────────────────────┬─────────────────────────────────────────┘
                            │ HTTP Requests (axios)
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                       BACKEND PROXY                                  │
│                  Python FastAPI (Port 8001)                          │
│   ┌─────────────────────────────────────────────────────────────┐  │
│   │ server.py                                                    │  │
│   │ - File upload handling (/api/data/upload)                   │  │
│   │ - File mapping management (/api/data/process)               │  │
│   │ - R Engine proxy (/api/r/*)                                 │  │
│   │ - R Engine status & control                                 │  │
│   └─────────────────────────────────────────────────────────────┘  │
└───────────────────────────┬─────────────────────────────────────────┘
                            │ HTTP Proxy (httpx)
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                      ANALYTICS ENGINE                                │
│                   R Plumber API (Port 8002)                          │
│   ┌─────────────────────────────────────────────────────────────┐  │
│   │ plumber_api.R                                                │  │
│   │ - Load data files (load_data_file function)                 │  │
│   │ - Analytics endpoints (/analytics/*)                        │  │
│   │ - Optimizer endpoint (/optimizer/run)                       │  │
│   │ - Health & status checks                                    │  │
│   └─────────────────────────────────────────────────────────────┘  │
└───────────────────────────┬─────────────────────────────────────────┘
                            │ File I/O (readxl, data.table)
                            ▼
┌─────────────────────────────────────────────────────────────────────┐
│                       DATA STORAGE                                   │
│                    /app/data Directory                               │
│   ┌─────────────────────────────────────────────────────────────┐  │
│   │ - file_mapping.json (maps file_id to file paths)            │  │
│   │ - Uploaded Excel files (.xlsx)                              │  │
│   │ - Uploaded CSV files (.csv)                                 │  │
│   └─────────────────────────────────────────────────────────────┘  │
└─────────────────────────────────────────────────────────────────────┘
'''

doc.add_paragraph(arch_diagram)

doc.add_page_break()

# Section 3: R Engine Files
doc.add_heading('3. R Engine Files - Detailed Description', level=1)

doc.add_paragraph(
    'All R files are located in: /app/backend/r_engine/'
)

doc.add_heading('3.1 Core API File', level=2)

# plumber_api.R
doc.add_heading('plumber_api.R (20 KB) - Main REST API', level=3)
doc.add_paragraph(
    'This is the PRIMARY file that exposes R functionality as REST endpoints. '
    'It is the entry point for all analytics requests from the frontend.'
)

p = doc.add_paragraph()
p.add_run('Key Functions:\n').bold = True
p.add_run('''
• load_file_mapping() - Loads /app/data/file_mapping.json to find uploaded files
• load_data_file(file_id) - Reads Excel/CSV files using readxl and data.table
• cached_data - In-memory cache for loaded datasets

Key Endpoints:
• GET /health - Returns R engine status and version
• GET /status - Returns working directory and loaded file info
• GET /analytics/filters - Returns filter options from Nielsen data (brands, ppgs, customers)
• GET /analytics/executive-summary - Returns brand-level aggregated data
• GET /analytics/brand-summary - Returns brand performance summary
• GET /analytics/roi-analysis - Returns ROI analysis from model results
• GET /analytics/events - Returns historical promotion events
• POST /optimizer/run - MAIN ENDPOINT: Runs optimization with constraints

Data Loading Process (lines 31-63):
1. Reads file_mapping.json to get file paths
2. Checks cache for previously loaded data
3. Uses readxl::read_excel for .xlsx files
4. Uses data.table::fread for .csv files
5. Caches loaded data for performance
''')

doc.add_heading('3.2 Startup File', level=2)

# start_api.R
doc.add_heading('start_api.R (794 bytes) - API Launcher', level=3)
doc.add_paragraph(
    'Simple startup script that initializes and runs the Plumber API on port 8002.'
)
p = doc.add_paragraph()
p.add_run('Code Flow:\n').bold = True
p.add_run('''
1. Load plumber library
2. Load plumber_api.R as the API definition
3. Call pr_run() to start server on port 8002
4. Print available endpoints for debugging
''')

doc.add_heading('3.3 Original R-Shiny Files (For Reference)', level=2)

doc.add_paragraph(
    'These files contain the original R-Shiny application logic. They are kept '
    'for reference and some functions may be called by plumber_api.R:'
)

r_files_table = doc.add_table(rows=12, cols=3)
r_files_table.style = 'Table Grid'

headers = ['File Name', 'Size', 'Purpose']
for i, h in enumerate(headers):
    r_files_table.rows[0].cells[i].text = h

r_files_data = [
    ('server.R', '912 KB', 'Main R-Shiny server logic - contains optimization algorithms, data processing, reactive functions'),
    ('global.R', '121 KB', 'Global variables, helper functions, plot generation code'),
    ('ui.R', '85 KB', 'R-Shiny UI definition (not used in React version)'),
    ('annual_optimization.R', '128 KB', 'Annual optimization algorithms - constraint handling, promo sequencing'),
    ('data_prep_event_list.R', '125 KB', 'Data preparation functions for event/promotion lists'),
    ('constraint_fun.R', '5.7 KB', 'Constraint validation functions (margin, spend, duration checks)'),
    ('do_calculation.R', '3 KB', 'Core calculation functions for ROI, revenue, volume'),
    ('Competitor_seq.R', '12.8 KB', 'Competitor analysis and sequencing logic'),
    ('Recalculate_Function.R', '2.8 KB', 'Functions to recalculate metrics when constraints change'),
    ('ppg_budget_check_fun.R', '3.3 KB', 'Budget validation functions per PPG'),
    ('pattern_detection.R', '2.3 KB', 'Pattern detection for promotional timing')
]

for i, (fname, size, purpose) in enumerate(r_files_data, 1):
    r_files_table.rows[i].cells[0].text = fname
    r_files_table.rows[i].cells[1].text = size
    r_files_table.rows[i].cells[2].text = purpose

doc.add_page_break()

# Section 4: Data Upload Flow
doc.add_heading('4. Data Upload & Processing Flow', level=1)

doc.add_heading('4.1 Upload Process', level=2)

upload_flow = '''
STEP-BY-STEP DATA UPLOAD FLOW:

1. USER ACTION: Click Settings button → Open File Manager
   Frontend File: /app/frontend/src/pages/DataInputPage.js

2. FILE SELECTION: User selects Excel/CSV file for a data category
   Categories defined in REQUIRED_FILES array (lines 28-100):
   - nielsen_rms: Nielsen RMS Data (sales, volume, value)
   - model_results: Elasticity model outputs
   - events: Historical promotion events
   - cost_bible: Product cost information
   - retailer_slots: Retailer promotional calendar
   - ean_ldesc: EAN to product mapping
   - lsm: Last Sold Month data
   - brand_delist: Brand delist/new list
   - retailer_weekend: Retailer week configuration

3. UPLOAD REQUEST: Frontend sends POST to /api/data/upload
   Backend File: /app/backend/server.py (lines 82-126)
   
   Request format:
   - file: The uploaded file (multipart/form-data)
   - file_id: Category identifier (e.g., "nielsen_rms")
   
   Backend actions:
   a. Validates file extension (.xlsx, .xls, .csv)
   b. Saves file to /app/data/{filename}
   c. Updates file_mapping.json with file_id → path mapping

4. FILE MAPPING UPDATE: Backend updates /app/data/file_mapping.json
   Example content:
   {
     "nielsen_rms": "/app/data/New file.xlsx",
     "model_results": "/app/data/Model Results for 5 SKUs.xlsx",
     "events": "/app/data/HEA Events.xlsx"
   }

5. R ENGINE READS DATA: When optimization runs, R reads file_mapping.json
   R File: /app/backend/r_engine/plumber_api.R
   
   Function: load_data_file(file_id)
   - Reads file_mapping.json
   - Gets file path for requested file_id
   - Loads Excel using readxl::read_excel
   - Caches data in memory for subsequent requests
'''

doc.add_paragraph(upload_flow)

doc.add_page_break()

doc.add_heading('4.2 File Mapping JSON Structure', level=2)
doc.add_paragraph(
    'Location: /app/data/file_mapping.json\n\n'
    'This file maps logical file IDs to physical file paths. The R engine reads '
    'this file to know where each data source is located.'
)

mapping_example = '''
{
  "nielsen_rms": "/app/data/New file.xlsx",
  "model_results": "/app/data/Model Results for 5 SKUs.xlsx",
  "events": "/app/data/HEA Events.xlsx",
  "cost_bible": "/app/data/Updated Cost Bible.xlsx",
  "retailer_slots": "/app/data/Carrefour Slots 2025-2026.xlsx",
  "ean_ldesc": "/app/data/8 Nielsen EAN to LDESC D77.xlsx",
  "lsm": "/app/data/LSM_New.xlsx"
}
'''
doc.add_paragraph(mapping_example)

doc.add_page_break()

# Section 5: Frontend React Components
doc.add_heading('5. Frontend React Components', level=1)

doc.add_heading('5.1 Main Optimizer Page', level=2)
doc.add_paragraph('File: /app/frontend/src/pages/OptimizerNewPage.js')

p = doc.add_paragraph()
p.add_run('Component Structure:\n').bold = True
p.add_run('''
1. STATE MANAGEMENT (lines 97-160):
   - config: Global constraints (goal, margins, spend, max weeks)
   - productConstraints: Product-level constraints (investment, price, mechanics)
   - promoData: Array of optimized promotions (displayed in calendar/table)
   - summaryData: KPI summary (revenue, ROI, margin, etc.)
   - rEngineStatus: R engine connection status
   - filters: Dropdown options from R API

2. R ENGINE STATUS CHECK (lines 170-210):
   - checkREngineStatus(): Calls /api/r-engine/status
   - startREngine(): Calls /api/r-engine/start
   - Auto-connect on page load
   - Periodic check every 30 seconds

3. FILTERS LOADING (lines 227-243):
   - Calls GET /api/promo/filters
   - Populates Retailer and PPG dropdowns
   - Uses real data from Nielsen file

4. OPTIMIZATION EXECUTION (lines 245-330):
   handleRunOptimizer():
   a. Builds URL parameters from constraints
   b. POST to /api/r/optimizer/run?{params}
   c. Receives promos array and summary from R
   d. Transforms R JSON format (arrays) to JS objects
   e. Updates promoData and summaryData state
   f. Updates calendar and table display
''')

doc.add_heading('5.2 Data Input Page', level=2)
doc.add_paragraph('File: /app/frontend/src/pages/DataInputPage.js')

p = doc.add_paragraph()
p.add_run('Component Functions:\n').bold = True
p.add_run('''
1. REQUIRED_FILES constant (lines 28-100):
   - Defines all data file categories
   - Each has: id, name, description, defaultName, required, category

2. File Upload Handler:
   - Uses axios to POST /api/data/upload
   - Sends file as FormData with file_id
   - Shows success/error toast notifications

3. File Status Display:
   - Shows which files are uploaded
   - Green checkmark for uploaded files
   - Red warning for missing required files

4. Process Button:
   - POST /api/data/process with file mapping
   - R engine reloads all data files
''')

doc.add_page_break()

# Section 6: API Endpoints
doc.add_heading('6. API Endpoints & Data Flow', level=1)

doc.add_heading('6.1 FastAPI Endpoints (server.py)', level=2)

api_table = doc.add_table(rows=10, cols=3)
api_table.style = 'Table Grid'

api_headers = ['Endpoint', 'Method', 'Description']
for i, h in enumerate(api_headers):
    api_table.rows[0].cells[i].text = h

api_data = [
    ('/api/health', 'GET', 'Health check - returns FastAPI and R engine status'),
    ('/api/data/upload', 'POST', 'Upload Excel/CSV file with file_id mapping'),
    ('/api/data/files', 'GET', 'List all uploaded files in /app/data'),
    ('/api/data/files/{name}', 'DELETE', 'Delete a specific uploaded file'),
    ('/api/data/process', 'POST', 'Save file mapping and trigger R reload'),
    ('/api/r-engine/status', 'GET', 'Get R engine connection status'),
    ('/api/r-engine/start', 'POST', 'Start the R Plumber API process'),
    ('/api/r/{path}', 'GET/POST', 'Proxy any request to R API at localhost:8002'),
    ('/api/promo/filters', 'GET', 'Get filter options (proxied to R)')
]

for i, (endpoint, method, desc) in enumerate(api_data, 1):
    api_table.rows[i].cells[0].text = endpoint
    api_table.rows[i].cells[1].text = method
    api_table.rows[i].cells[2].text = desc

doc.add_paragraph()

doc.add_heading('6.2 R Plumber Endpoints (plumber_api.R)', level=2)

r_api_table = doc.add_table(rows=9, cols=3)
r_api_table.style = 'Table Grid'

for i, h in enumerate(api_headers):
    r_api_table.rows[0].cells[i].text = h

r_api_data = [
    ('/health', 'GET', 'R engine health - returns status and R version'),
    ('/status', 'GET', 'System status - working dir, loaded files'),
    ('/analytics/filters', 'GET', 'Filter options from Nielsen data'),
    ('/analytics/executive-summary', 'GET', 'Brand aggregations from Nielsen'),
    ('/analytics/brand-summary', 'GET', 'Brand performance metrics'),
    ('/analytics/roi-analysis', 'GET', 'ROI data from model results'),
    ('/analytics/events', 'GET', 'Historical promotion events'),
    ('/optimizer/run', 'POST', 'MAIN: Run optimization with constraints')
]

for i, (endpoint, method, desc) in enumerate(r_api_data, 1):
    r_api_table.rows[i].cells[0].text = endpoint
    r_api_table.rows[i].cells[1].text = method
    r_api_table.rows[i].cells[2].text = desc

doc.add_page_break()

# Section 7: Optimization Calculation Logic
doc.add_heading('7. Optimization Calculation Logic', level=1)

doc.add_heading('7.1 Optimizer Endpoint: /optimizer/run', level=2)
doc.add_paragraph('File: /app/backend/r_engine/plumber_api.R (lines 408-580)')

p = doc.add_paragraph()
p.add_run('Input Parameters:\n').bold = True
p.add_run('''
• goal: "revenue" | "volume" | "margin" | "spend" | "roi"
• gm_min, gm_max: Gross margin range (%)
• spend_min, spend_max: Total trade spend limits ($)
• max_promos: Maximum promotional periods
• investment_min, investment_max: Per-promo investment range ($)
• price_min, price_max: Promo price range ($)
• allowed_mechanics: JSON array ["TPR", "Display", "Flyer", "Display + Flyer"]
• allow_consecutive: Boolean - allow back-to-back promos
''')

doc.add_heading('7.2 Calculation Steps', level=2)

calc_steps = '''
STEP 1: Load Data Files
-----------------------
• events_data = load_data_file("events")    # Historical promos
• model_data = load_data_file("model_results")  # Elasticity coefficients
• nielsen_data = load_data_file("nielsen_rms")  # Sales data

STEP 2: Calculate Base Values from Nielsen
------------------------------------------
• total_value = sum(nielsen_data$Value)     # Total sales value
• total_units = sum(nielsen_data$Units)     # Total units sold
• base_price = total_value / total_units    # Average price
• base_volume = total_value / 52 / base_price  # Weekly volume estimate

STEP 3: Process Events Data
---------------------------
• Add month column from Date field
• Determine mechanic type from flags:
  - Flag_TPR_HEA == 1 && Flag_Display_HEA == 1 → "Display + Flyer"
  - Flag_Display_HEA == 1 → "Display"
  - Flag_TPR_HEA == 1 → "TPR"
  - Otherwise → "Flyer"
• Filter by allowed_mechanics parameter
• Aggregate by month + mechanic

STEP 4: Sort by Optimization Goal
---------------------------------
• If goal == "revenue": Sort by Weekly Cost (descending)
• If goal == "roi": Sort by Weekly Cost / Net Investment (descending)
• If goal == "volume": Sort by volume potential (descending)

STEP 5: Generate Optimized Promos
---------------------------------
For each month in sorted data:
  a. Check consecutive constraint
  b. Calculate investment within min/max bounds
  c. Check if total investment exceeds spend_max
  d. Apply volume multiplier based on goal:
     - revenue: 1.2 - 1.8x
     - volume: 1.5 - 2.0x
     - roi: 0.8 - 1.2x
  e. Calculate revenue = volume × price
  f. Calculate ROI = (revenue - investment) / investment × 100
  g. Add promo to list with:
     - slot, month, row (calendar position)
     - type (mechanic)
     - price, tpr, investment
     - volume, revenue, gm, roi
     - startDay, endDay

STEP 6: Calculate Summary Metrics
---------------------------------
• totalRevenue = sum of all promo revenues
• totalInvestment = sum of all promo investments
• avgGrossMargin = mean of all promo GMs
• baselineRevenue = totalRevenue × 0.75
• incrementalRevenue = totalRevenue - baselineRevenue
• roi = (incrementalRevenue / totalInvestment) × 100
• constraintsMet = check all constraints satisfied

STEP 7: Return Result
---------------------
Return JSON with:
• success: TRUE/FALSE
• data_source: "events_data" or "nielsen_data"
• constraints_met: TRUE/FALSE
• promos: Array of promo objects
• summary: KPI summary object
'''

doc.add_paragraph(calc_steps)

doc.add_page_break()

# Section 8: Configuration Parameters
doc.add_heading('8. Configuration Parameters', level=1)

doc.add_heading('8.1 Global Constraints', level=2)

gc_table = doc.add_table(rows=4, cols=4)
gc_table.style = 'Table Grid'

gc_headers = ['Parameter', 'Variable', 'Default', 'Used In']
for i, h in enumerate(gc_headers):
    gc_table.rows[0].cells[i].text = h

gc_data = [
    ('Gross Margin Range', 'gmRangeMin/gmRangeMax', '28% - 35%', 'R optimizer constraint check'),
    ('Trade Spend Range', 'tradeSpendMin/tradeSpendMax', '$50K - $200K', 'R optimizer loop limit'),
    ('Max Promo Weeks', 'maxPromoWeeks', '12', 'R optimizer promo count limit')
]

for i, (param, var, default, used) in enumerate(gc_data, 1):
    gc_table.rows[i].cells[0].text = param
    gc_table.rows[i].cells[1].text = var
    gc_table.rows[i].cells[2].text = default
    gc_table.rows[i].cells[3].text = used

doc.add_paragraph()

doc.add_heading('8.2 Product Constraints', level=2)

pc_table = doc.add_table(rows=6, cols=4)
pc_table.style = 'Table Grid'

for i, h in enumerate(gc_headers):
    pc_table.rows[0].cells[i].text = h

pc_data = [
    ('Investment per Promo', 'investmentMin/investmentMax', '$2K - $15K', 'R optimizer per-promo limit'),
    ('Promo Price Range', 'priceMin/priceMax', '$3.49 - $4.49', 'R optimizer price generation'),
    ('Allowed Mechanics', 'allowedMechanics object', 'TPR, Display, Flyer', 'R optimizer mechanic filter'),
    ('Promo Duration', 'durationMin/durationMax', '1 - 2 weeks', 'Calendar date range'),
    ('Allow Consecutive', 'allowConsecutive', 'true', 'R optimizer skip logic')
]

for i, (param, var, default, used) in enumerate(pc_data, 1):
    pc_table.rows[i].cells[0].text = param
    pc_table.rows[i].cells[1].text = var
    pc_table.rows[i].cells[2].text = default
    pc_table.rows[i].cells[3].text = used

doc.add_page_break()

# Section 9: Output Rendering
doc.add_heading('9. Output & Results Rendering', level=1)

doc.add_heading('9.1 Data Flow: R → React', level=2)

render_flow = '''
R API RESPONSE FORMAT:
----------------------
{
  "success": [true],
  "data_source": ["events_data"],
  "constraints_met": [true],
  "promos": [
    {
      "slot": [1],
      "month": [0],
      "row": [1],
      "type": ["Flyer"],
      "price": [3.87],
      "tpr": [24.5],
      "investment": [2000],
      "volume": [25559],
      "revenue": [98875],
      "gm": [29.4],
      "roi": [4843.7],
      "startDay": [11],
      "endDay": [20]
    },
    ...
  ],
  "summary": {
    "totalRevenue": [541395],
    "incrementalRevenue": [135349],
    "tradeSpend": [24000],
    "roi": [564.0],
    "avgGrossMargin": [31.2],
    "totalPromos": [12],
    "baselineRevenue": [406046],
    "revenueVsBaseline": [33.3]
  }
}

REACT TRANSFORMATION (OptimizerNewPage.js lines 285-320):
---------------------------------------------------------
// R returns arrays even for single values, need to extract [0]
const promos = result.promos.map(p => ({
  slot: Array.isArray(p.slot) ? p.slot[0] : p.slot,
  month: Array.isArray(p.month) ? p.month[0] : p.month,
  // ... same for all fields
}));

const extractValue = (val) => Array.isArray(val) ? val[0] : val;

setSummaryData({
  totalRevenue: extractValue(summary.totalRevenue) || 0,
  // ... same for all summary fields
});

setPromoData(promos);  // Triggers re-render of calendar & table
'''

doc.add_paragraph(render_flow)

doc.add_heading('9.2 UI Components That Render Data', level=2)

ui_components = '''
1. KPI CARDS (lines 680-780):
   - Display summaryData.totalRevenue, incrementalRevenue, tradeSpend, roi
   - Format using formatCurrency() helper
   - Show vs baseline percentage

2. PROMOTION CALENDAR (lines 810-900):
   - 3 rows × 12 columns grid
   - Row 0 = 1st of month, Row 1 = 11th, Row 2 = 21st
   - getPromoForCell(monthIdx, rowIdx) finds matching promo
   - Color-coded by type using PROMO_COLORS constant
   - Tooltip shows full promo details on hover

3. PROMOTION DETAILS TABLE (lines 920-1000):
   - Maps over promoData array
   - Columns: SLOT, PERIOD, MECHANIC, PRICE, TPR%, INVESTMENT, INCR.VOL, REVENUE, GM%, ROI%
   - Sortable by slot
   - Color-coded mechanic badges
'''

doc.add_paragraph(ui_components)

doc.add_page_break()

# Section 10: File Directory Structure
doc.add_heading('10. File Directory Structure', level=1)

dir_structure = '''
/app/
├── frontend/
│   ├── src/
│   │   ├── pages/
│   │   │   ├── OptimizerNewPage.js    # Main optimizer UI (standalone)
│   │   │   ├── DataInputPage.js       # File upload interface
│   │   │   ├── PromoAnalysisPage.js   # Executive summary dashboard
│   │   │   └── ...
│   │   ├── components/
│   │   │   ├── ui/                    # shadcn/ui components
│   │   │   └── layout/
│   │   │       ├── MainLayout.js      # Layout wrapper with sidebar
│   │   │       └── Sidebar.js         # Navigation sidebar
│   │   └── App.js                     # Route definitions
│   ├── public/
│   │   ├── reckitt_logo.svg           # Brand logo
│   │   └── Trade_Promotion_Optimizer_Documentation.docx
│   └── .env                           # REACT_APP_BACKEND_URL
│
├── backend/
│   ├── server.py                      # FastAPI main application
│   ├── config.py                      # Environment configuration
│   ├── routes/
│   │   └── promo.py                   # Promo-related endpoints
│   ├── r_engine/
│   │   ├── plumber_api.R              # REST API endpoints (MAIN)
│   │   ├── start_api.R                # API startup script
│   │   ├── server.R                   # Original Shiny server logic
│   │   ├── global.R                   # Helper functions
│   │   ├── annual_optimization.R      # Optimization algorithms
│   │   ├── data_prep_event_list.R     # Data preparation
│   │   ├── constraint_fun.R           # Constraint validation
│   │   └── ...
│   ├── requirements.txt               # Python dependencies
│   └── .env                           # MONGO_URL, DB_NAME (not used)
│
├── data/                              # Uploaded data files
│   ├── file_mapping.json              # File ID to path mapping
│   ├── New file.xlsx                  # Nielsen RMS data
│   ├── Model Results for 5 SKUs.xlsx  # Model results
│   ├── HEA Events.xlsx                # Events data
│   └── ...
│
└── memory/
    └── PRD.md                         # Product requirements document
'''

doc.add_paragraph(dir_structure)

doc.add_page_break()

# Section 11: Troubleshooting
doc.add_heading('11. Troubleshooting Guide', level=1)

troubleshooting = [
    ('R Engine Not Starting', 
     'Check if R is installed: which Rscript\n'
     'Install if missing: apt-get install r-base\n'
     'Install R packages: R -e "install.packages(c(\'plumber\', \'jsonlite\', \'data.table\', \'readxl\'))"'),
    
    ('Port 8002 Already in Use',
     'Kill existing process: fuser -k 8002/tcp\n'
     'Wait 2 seconds, then restart R API'),
    
    ('Data Not Loading in R',
     'Check file_mapping.json exists in /app/data/\n'
     'Verify file paths in mapping are correct\n'
     'Check R logs: tail -50 /var/log/r_api.log'),
    
    ('Optimization Returns Empty Results',
     'Ensure at least one mechanic is selected\n'
     'Check that uploaded files have data\n'
     'Verify constraints are not too restrictive'),
    
    ('Frontend Shows "R Engine Stopped"',
     'Click the power button to start R Engine\n'
     'Check backend proxy is running: sudo supervisorctl status\n'
     'Verify R API is accessible: curl http://localhost:8002/health')
]

for issue, solution in troubleshooting:
    doc.add_heading(issue, level=2)
    doc.add_paragraph(solution)

doc.add_page_break()

# Section 12: Appendix
doc.add_heading('12. Appendix', level=1)

doc.add_heading('12.1 Glossary', level=2)

glossary = [
    ('FMCG', 'Fast-Moving Consumer Goods'),
    ('TPR', 'Temporary Price Reduction - a promotional mechanic'),
    ('PPG', 'Product Performance Group'),
    ('ROI', 'Return on Investment'),
    ('GM', 'Gross Margin'),
    ('Nielsen RMS', 'Nielsen Retail Measurement Services - point of sale data'),
    ('KPI', 'Key Performance Indicator'),
    ('Plumber', 'R package for creating REST APIs from R functions'),
    ('FastAPI', 'Modern Python web framework for building APIs')
]

for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(f'{term}: ').bold = True
    p.add_run(definition)

doc.add_heading('12.2 Key Code References', level=2)

code_refs = '''
FRONTEND:
• Optimizer Page: /app/frontend/src/pages/OptimizerNewPage.js
• Data Upload: /app/frontend/src/pages/DataInputPage.js
• API Base URL: process.env.REACT_APP_BACKEND_URL

BACKEND:
• FastAPI Server: /app/backend/server.py
• R Engine Proxy: Lines 280-310 in server.py
• File Upload: Lines 82-126 in server.py

R ENGINE:
• Main API: /app/backend/r_engine/plumber_api.R
• Optimizer Function: Lines 408-580 in plumber_api.R
• Data Loading: load_data_file() function, lines 31-63
'''

doc.add_paragraph(code_refs)

# Save document
doc.save('/app/frontend/public/Trade_Promotion_Optimizer_Technical_Doc.docx')
print('Technical document created successfully!')
