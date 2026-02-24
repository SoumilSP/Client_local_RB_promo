from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

# Create document
doc = Document()

# Set document title
title = doc.add_heading('Trade Promotion Optimizer', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('FMCG Promo Calendar Optimization Engine')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(196, 25, 116)  # Reckitt pink

# Add date
date_para = doc.add_paragraph(f'Document Version: 1.0 | Date: {datetime.now().strftime("%B %d, %Y")}')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(10)
date_para.runs[0].font.italic = True

doc.add_paragraph()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Overview',
    '3. Key Features',
    '4. User Interface Guide',
    '5. Configuration Parameters',
    '6. Data Sources & Integration',
    '7. Optimization Goals',
    '8. Output & Results',
    '9. Technical Architecture',
    '10. Appendix'
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# Section 1: Executive Summary
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'The Trade Promotion Optimizer is a sophisticated FMCG (Fast-Moving Consumer Goods) '
    'promotional planning tool designed to maximize return on trade investments. Built on '
    'a modern React frontend with an R-powered analytics engine, the system enables '
    'category managers and trade marketing teams to create optimized promotional calendars '
    'that balance business objectives with operational constraints.'
)

doc.add_heading('Key Benefits', level=2)
benefits = [
    'Data-Driven Decision Making: Leverages historical promotion data and Nielsen RMS data for accurate forecasting',
    'Constraint-Based Optimization: Ensures all promotions meet business rules and financial guardrails',
    'Real-Time Results: Instant optimization with visual calendar output',
    'Flexible Goal Setting: Supports multiple optimization objectives (Revenue, Volume, Margin, ROI)',
    'Interactive Configuration: Intuitive UI for adjusting constraints and re-running scenarios'
]
for benefit in benefits:
    doc.add_paragraph(benefit, style='List Bullet')

doc.add_page_break()

# Section 2: System Overview
doc.add_heading('2. System Overview', level=1)

doc.add_heading('2.1 Architecture', level=2)
doc.add_paragraph(
    'The Trade Promotion Optimizer follows a three-tier architecture:'
)

# Architecture table
arch_table = doc.add_table(rows=4, cols=3)
arch_table.style = 'Table Grid'
arch_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
header_cells = arch_table.rows[0].cells
header_cells[0].text = 'Layer'
header_cells[1].text = 'Technology'
header_cells[2].text = 'Purpose'

# Data rows
data = [
    ('Frontend', 'React + Tailwind CSS', 'User interface, visualization, configuration'),
    ('Backend Proxy', 'Python FastAPI', 'API routing, file management, data proxy'),
    ('Analytics Engine', 'R Plumber API', 'Optimization algorithms, data processing')
]

for i, (layer, tech, purpose) in enumerate(data, 1):
    arch_table.rows[i].cells[0].text = layer
    arch_table.rows[i].cells[1].text = tech
    arch_table.rows[i].cells[2].text = purpose

doc.add_paragraph()

doc.add_heading('2.2 Data Flow', level=2)
doc.add_paragraph(
    '1. User uploads Excel/CSV data files (Nielsen RMS, Events, Model Results)\n'
    '2. Files are stored in the data directory and mapped for R engine access\n'
    '3. User configures optimization parameters via the UI\n'
    '4. Frontend sends optimization request to FastAPI backend\n'
    '5. Backend proxies request to R Plumber API with all constraints\n'
    '6. R engine processes data and returns optimized promo calendar\n'
    '7. Results displayed in interactive calendar and detailed table view'
)

doc.add_page_break()

# Section 3: Key Features
doc.add_heading('3. Key Features', level=1)

doc.add_heading('3.1 Optimizer Configuration Panel', level=2)
doc.add_paragraph(
    'The left-side configuration panel provides comprehensive control over optimization parameters:'
)

features = [
    ('Context Selection', 'Choose Retailer and PPG (Product Performance Group) to focus optimization'),
    ('Optimization Goal', 'Select from 5 optimization objectives: Maximize Revenue, Maximize Volume, Maximize Margin, Minimize Spend, or Maximize ROI'),
    ('Global Constraints', 'Set portfolio-level boundaries including Gross Margin Range, Trade Spend Range, and Maximum Promo Weeks per Year'),
    ('Product Constraints', 'Define promotion-level limits including Investment per Promo, Price Range, Allowed Mechanics, Duration, and Consecutive Promo Rules')
]

for feature, desc in features:
    p = doc.add_paragraph()
    p.add_run(f'{feature}: ').bold = True
    p.add_run(desc)

doc.add_heading('3.2 Results Dashboard', level=2)
doc.add_paragraph('The right-side panel displays optimization results:')

results_features = [
    'KPI Summary Cards: Total Revenue, Incremental Revenue, Trade Spend, ROI, Avg Gross Margin, Total Promos, Baseline Revenue',
    'Promotion Calendar: Visual 12-month calendar with 3 time periods per month (1st, 11th, 21st)',
    'Color-Coded Mechanics: TPR (Purple), Display (Orange), Flyer (Cyan), Display + Flyer (Emerald)',
    'Interactive Tooltips: Hover over any promo to see detailed information',
    'Promotion Details Table: Sortable table with all promotion attributes'
]

for feature in results_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('3.3 R Engine Status Monitor', level=2)
doc.add_paragraph(
    'Real-time monitoring of the R analytics engine with:\n'
    '• Status indicator (Connected/Stopped/Starting)\n'
    '• One-click engine start button\n'
    '• Auto-connect on page load\n'
    '• Periodic health checks (every 30 seconds)'
)

doc.add_page_break()

# Section 4: User Interface Guide
doc.add_heading('4. User Interface Guide', level=1)

doc.add_heading('4.1 Header Bar', level=2)
header_items = [
    ('Reckitt Logo', 'Brand identification'),
    ('Title', 'Trade Promotion Optimizer'),
    ('R Engine Status', 'Shows connection status with start/refresh controls'),
    ('Download Button', 'Export results'),
    ('Settings Button', 'Access Data File Manager for uploading source files'),
    ('Help Button', 'Documentation and support')
]

for item, desc in header_items:
    p = doc.add_paragraph()
    p.add_run(f'{item}: ').bold = True
    p.add_run(desc)

doc.add_heading('4.2 Running an Optimization', level=2)
steps = [
    'Ensure R Engine status shows "Connected" (green)',
    'Select Retailer and PPG from the Context dropdowns',
    'Choose your Optimization Goal (Revenue, Volume, Margin, Spend, or ROI)',
    'Expand Global Constraints and set: Gross Margin Range (%), Trade Spend Range ($), Max Promo Weeks',
    'Expand Product Constraints and configure: Investment Range, Price Range, Allowed Mechanics, Duration, Consecutive Rules',
    'Click the "Run Optimizer" button',
    'Review results in the Annual Summary panel',
    'Hover over calendar cells for detailed promo information',
    'Scroll down to see the Promotion Details table'
]

for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_page_break()

# Section 5: Configuration Parameters
doc.add_heading('5. Configuration Parameters', level=1)

doc.add_heading('5.1 Global Constraints', level=2)

# Global constraints table
gc_table = doc.add_table(rows=4, cols=4)
gc_table.style = 'Table Grid'

gc_headers = ['Parameter', 'Description', 'Default Range', 'Unit']
for i, header in enumerate(gc_headers):
    gc_table.rows[0].cells[i].text = header

gc_data = [
    ('Gross Margin Range', 'Acceptable margin boundaries for promotions', '28% - 35%', 'Percentage'),
    ('Trade Spend Range', 'Total promotional investment limits', '$50,000 - $200,000', 'USD'),
    ('Max Promo Weeks/Year', 'Maximum promotional periods allowed', '12', 'Weeks')
]

for i, (param, desc, default, unit) in enumerate(gc_data, 1):
    gc_table.rows[i].cells[0].text = param
    gc_table.rows[i].cells[1].text = desc
    gc_table.rows[i].cells[2].text = default
    gc_table.rows[i].cells[3].text = unit

doc.add_paragraph()

doc.add_heading('5.2 Product Constraints', level=2)

# Product constraints table
pc_table = doc.add_table(rows=6, cols=4)
pc_table.style = 'Table Grid'

for i, header in enumerate(gc_headers):
    pc_table.rows[0].cells[i].text = header

pc_data = [
    ('Investment per Promo', 'Min/max spend per promotional event', '$2,000 - $15,000', 'USD'),
    ('Promo Price Range', 'Acceptable promoted price boundaries', '$3.49 - $4.49', 'USD'),
    ('Allowed Mechanics', 'Permitted promotion types', 'TPR, Display, Flyer', 'Checkbox'),
    ('Promo Duration', 'Length of promotional periods', '1 - 2', 'Weeks'),
    ('Allow Consecutive', 'Whether promos can run back-to-back', 'Yes', 'Boolean')
]

for i, (param, desc, default, unit) in enumerate(pc_data, 1):
    pc_table.rows[i].cells[0].text = param
    pc_table.rows[i].cells[1].text = desc
    pc_table.rows[i].cells[2].text = default
    pc_table.rows[i].cells[3].text = unit

doc.add_page_break()

# Section 6: Data Sources
doc.add_heading('6. Data Sources & Integration', level=1)

doc.add_heading('6.1 Required Data Files', level=2)

data_files = [
    ('Nielsen RMS Data', 'Historical sales data including volume, value, and price by brand/SKU', 'Excel (.xlsx)', 'Brand performance, baseline volume calculations'),
    ('Events Data', 'Historical promotion events with dates, mechanics, and costs', 'Excel (.xlsx)', 'Promo pattern analysis, mechanic distribution'),
    ('Model Results', 'Price elasticity and promotional lift coefficients', 'Excel (.xlsx)', 'ROI calculations, volume forecasting')
]

for name, desc, format_type, usage in data_files:
    p = doc.add_paragraph()
    p.add_run(f'{name}\n').bold = True
    p.add_run(f'Description: {desc}\n')
    p.add_run(f'Format: {format_type}\n')
    p.add_run(f'Usage: {usage}')
    doc.add_paragraph()

doc.add_heading('6.2 Data Processing', level=2)
doc.add_paragraph(
    'The R engine processes uploaded files to extract:\n'
    '• Historical promotion patterns by month and mechanic type\n'
    '• Baseline volume and revenue calculations from Nielsen data\n'
    '• Price elasticity for ROI projections\n'
    '• Investment efficiency metrics from past events'
)

doc.add_page_break()

# Section 7: Optimization Goals
doc.add_heading('7. Optimization Goals', level=1)

goals = [
    ('Maximize Revenue', 
     'Prioritizes promotions that generate the highest total revenue.',
     'Best for: Growth periods, market share battles, new product launches'),
    ('Maximize Volume', 
     'Focuses on unit sales and market penetration.',
     'Best for: Inventory clearance, distribution expansion, seasonal peaks'),
    ('Maximize Margin', 
     'Optimizes for profitability per unit sold.',
     'Best for: Premium positioning, profit improvement initiatives'),
    ('Minimize Spend', 
     'Reduces total promotional investment while maintaining effectiveness.',
     'Best for: Cost reduction programs, efficiency improvements'),
    ('Maximize ROI', 
     'Balances incremental revenue against promotional investment.',
     'Best for: Balanced portfolio management, sustainable growth')
]

for goal, desc, best_for in goals:
    doc.add_heading(goal, level=2)
    doc.add_paragraph(desc)
    p = doc.add_paragraph()
    p.add_run(best_for).italic = True

doc.add_page_break()

# Section 8: Output & Results
doc.add_heading('8. Output & Results', level=1)

doc.add_heading('8.1 KPI Metrics', level=2)

kpi_table = doc.add_table(rows=8, cols=2)
kpi_table.style = 'Table Grid'

kpi_headers = ['Metric', 'Description']
kpi_table.rows[0].cells[0].text = kpi_headers[0]
kpi_table.rows[0].cells[1].text = kpi_headers[1]

kpis = [
    ('Total Revenue', 'Sum of all promotional revenue across the calendar'),
    ('Incremental Revenue', 'Additional revenue generated above baseline'),
    ('Trade Spend', 'Total promotional investment'),
    ('ROI', 'Return on Investment: (Incremental Revenue / Trade Spend) × 100'),
    ('Avg Gross Margin', 'Weighted average margin across all promotions'),
    ('Total Promos', 'Number of promotional events scheduled'),
    ('Baseline Revenue', 'Expected revenue without promotions')
]

for i, (metric, desc) in enumerate(kpis, 1):
    kpi_table.rows[i].cells[0].text = metric
    kpi_table.rows[i].cells[1].text = desc

doc.add_paragraph()

doc.add_heading('8.2 Promotion Details Table Columns', level=2)

columns = [
    ('SLOT', 'Unique identifier for each promotional period'),
    ('PERIOD', 'Date range of the promotion'),
    ('MECHANIC', 'Type of promotion (TPR, Display, Flyer, or combination)'),
    ('PRICE', 'Promoted price point'),
    ('TPR %', 'Temporary Price Reduction percentage'),
    ('INVESTMENT', 'Cost of the promotional activity'),
    ('INCR. VOL', 'Incremental volume expected'),
    ('REVENUE', 'Expected promotional revenue'),
    ('GM %', 'Gross Margin percentage'),
    ('ROI %', 'Return on Investment for this promotion')
]

for col, desc in columns:
    p = doc.add_paragraph()
    p.add_run(f'{col}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# Section 9: Technical Architecture
doc.add_heading('9. Technical Architecture', level=1)

doc.add_heading('9.1 Technology Stack', level=2)

tech_table = doc.add_table(rows=6, cols=3)
tech_table.style = 'Table Grid'

tech_headers = ['Component', 'Technology', 'Version']
for i, header in enumerate(tech_headers):
    tech_table.rows[0].cells[i].text = header

tech_data = [
    ('Frontend Framework', 'React', '18.x'),
    ('UI Components', 'shadcn/ui + Tailwind CSS', '3.x'),
    ('Backend Proxy', 'Python FastAPI', '0.100+'),
    ('Analytics Engine', 'R with Plumber', '4.2.2'),
    ('Charts', 'Recharts', '2.x')
]

for i, (comp, tech, ver) in enumerate(tech_data, 1):
    tech_table.rows[i].cells[0].text = comp
    tech_table.rows[i].cells[1].text = tech
    tech_table.rows[i].cells[2].text = ver

doc.add_paragraph()

doc.add_heading('9.2 API Endpoints', level=2)

api_endpoints = [
    ('GET /api/health', 'System health check including R engine status'),
    ('GET /api/r-engine/status', 'Detailed R engine connection status'),
    ('POST /api/r-engine/start', 'Start the R analytics engine'),
    ('POST /api/r/optimizer/run', 'Execute optimization with constraints'),
    ('GET /api/promo/filters', 'Get available filter options (retailers, PPGs)')
]

for endpoint, desc in api_endpoints:
    p = doc.add_paragraph()
    p.add_run(endpoint).bold = True
    p.add_run(f'\n{desc}')

doc.add_page_break()

# Section 10: Appendix
doc.add_heading('10. Appendix', level=1)

doc.add_heading('10.1 Glossary', level=2)

glossary = [
    ('FMCG', 'Fast-Moving Consumer Goods'),
    ('TPR', 'Temporary Price Reduction'),
    ('PPG', 'Product Performance Group'),
    ('ROI', 'Return on Investment'),
    ('GM', 'Gross Margin'),
    ('Nielsen RMS', 'Nielsen Retail Measurement Services'),
    ('KPI', 'Key Performance Indicator')
]

for term, definition in glossary:
    p = doc.add_paragraph()
    p.add_run(f'{term}: ').bold = True
    p.add_run(definition)

doc.add_heading('10.2 Troubleshooting', level=2)

issues = [
    ('R Engine Stopped', 'Click the power button next to the status indicator to start the engine. Wait 5-10 seconds for initialization.'),
    ('No Data Displayed', 'Ensure data files are uploaded via Settings > Data File Manager. Verify R engine is connected.'),
    ('Constraints Not Met', 'Review constraint ranges. The system may not find valid solutions if constraints are too restrictive.'),
    ('Optimization Taking Long', 'Large datasets may require additional processing time. Check R engine logs for progress.')
]

for issue, solution in issues:
    p = doc.add_paragraph()
    p.add_run(f'Issue: {issue}\n').bold = True
    p.add_run(f'Solution: {solution}')
    doc.add_paragraph()

doc.add_heading('10.3 Contact & Support', level=2)
doc.add_paragraph(
    'For technical support or feature requests, please contact:\n'
    '• Technical Team: [Your Technical Team Email]\n'
    '• Product Owner: [Product Owner Email]\n'
    '• Documentation: [Documentation URL]'
)

# Save document
doc.save('/app/frontend/public/Trade_Promotion_Optimizer_Documentation.docx')
print('Document created successfully!')
